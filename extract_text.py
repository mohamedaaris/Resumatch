"""
Text extraction module for ResuMatch AI
Handles OCR and document parsing for PDF, DOCX, and image files
"""

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import docx
import io
import logging
from typing import Optional, Union
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TextExtractor:
    """Handles text extraction from various document formats"""
    
    def __init__(self):
        """Initialize the text extractor"""
        # Set tesseract path for Windows (adjust if needed)
        if os.name == 'nt':  # Windows
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text
        """
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += page_text + "\n"
            
            doc.close()
            logger.info(f"Successfully extracted text from PDF ({len(text)} characters)")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX ({len(text)} characters)")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_from_image(self, file_path: str) -> str:
        """
        Extract text from image using OCR (Tesseract)
        
        Args:
            file_path (str): Path to the image file
            
        Returns:
            str: Extracted text
        """
        try:
            logger.info(f"Extracting text from image: {file_path}")
            
            # Open image with PIL
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use pytesseract for OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            logger.info(f"Successfully extracted text from image ({len(text)} characters)")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise
    
    def extract_from_pdf_with_ocr(self, file_path: str) -> str:
        """
        Extract text from PDF using OCR (for scanned PDFs)
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text
        """
        try:
            logger.info(f"Extracting text from PDF using OCR: {file_path}")
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Convert page to image
                mat = fitz.Matrix(2.0, 2.0)  # Increase resolution
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to PIL Image
                image = Image.open(io.BytesIO(img_data))
                
                # Use OCR
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += page_text + "\n"
            
            doc.close()
            logger.info(f"Successfully extracted text from PDF using OCR ({len(text)} characters)")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF using OCR: {str(e)}")
            raise
    
    def extract_text(self, file_path: str, file_type: str = None) -> str:
        """
        Main method to extract text from any supported file format
        
        Args:
            file_path (str): Path to the file
            file_type (str, optional): File type hint. If None, will be inferred from extension
            
        Returns:
            str: Extracted text
        """
        try:
            # Determine file type if not provided
            if file_type is None:
                file_type = file_path.lower().split('.')[-1]
            
            logger.info(f"Extracting text from {file_type.upper()} file: {file_path}")
            
            if file_type == 'pdf':
                # Try regular PDF extraction first
                try:
                    return self.extract_from_pdf(file_path)
                except:
                    # Fall back to OCR if regular extraction fails
                    logger.info("Regular PDF extraction failed, trying OCR...")
                    return self.extract_from_pdf_with_ocr(file_path)
            
            elif file_type == 'docx':
                return self.extract_from_docx(file_path)
            
            elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif']:
                return self.extract_from_image(file_path)
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error in extract_text: {str(e)}")
            raise
    
    def extract_text_from_bytes(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from file bytes (useful for Flask uploads)
        
        Args:
            file_bytes (bytes): File content as bytes
            file_type (str): File type (pdf, docx, jpg, etc.)
            
        Returns:
            str: Extracted text
        """
        try:
            logger.info(f"Extracting text from {file_type.upper()} bytes")
            
            if file_type == 'pdf':
                # Try regular PDF extraction first
                try:
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    text = ""
                    
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        text += page_text + "\n"
                    
                    doc.close()
                    return text
                except:
                    # Fall back to OCR
                    logger.info("Regular PDF extraction failed, trying OCR...")
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    text = ""
                    
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        mat = fitz.Matrix(2.0, 2.0)
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        page_text = pytesseract.image_to_string(image, lang='eng')
                        text += page_text + "\n"
                    
                    doc.close()
                    return text
            
            elif file_type == 'docx':
                doc = docx.Document(io.BytesIO(file_bytes))
                text = ""
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                
                return text
            
            elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif']:
                image = Image.open(io.BytesIO(file_bytes))
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                return pytesseract.image_to_string(image, lang='eng')
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error in extract_text_from_bytes: {str(e)}")
            raise


# Convenience function for easy usage
def extract_text_from_file(file_path: str, file_type: str = None) -> str:
    """
    Convenience function to extract text from a file
    
    Args:
        file_path (str): Path to the file
        file_type (str, optional): File type hint
        
    Returns:
        str: Extracted text
    """
    extractor = TextExtractor()
    return extractor.extract_text(file_path, file_type)


def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
    """
    Convenience function to extract text from file bytes
    
    Args:
        file_bytes (bytes): File content as bytes
        file_type (str): File type
        
    Returns:
        str: Extracted text
    """
    extractor = TextExtractor()
    return extractor.extract_text_from_bytes(file_bytes, file_type)


def extract_text_with_hints_from_bytes(file_bytes: bytes, file_type: str) -> dict:
    """
    Extract text and style/structural hints from file bytes.
    Returns a dict: { 'text': str, 'hints': { 'headings': [...], 'sections': [...], 'bold_lines': [...], 'education_candidates': [...] } }
    """
    extractor = TextExtractor()
    text = extractor.extract_text_from_bytes(file_bytes, file_type)
    hints = {
        'headings': [],
        'sections': [],
        'bold_lines': [],
        'education_candidates': []
    }
    try:
        if file_type == 'docx':
            doc = docx.Document(io.BytesIO(file_bytes))
            current_section = None
            sections_map = {}
            # Helper to normalize heading text
            def norm_heading(s: str) -> str:
                return (s or '').strip().replace(':','').upper()
            for p in doc.paragraphs:
                pt = p.text.strip()
                if not pt:
                    continue
                is_heading = False
                try:
                    if p.style and p.style.name and p.style.name.lower().startswith('heading'):
                        is_heading = True
                except Exception:
                    pass
                # Fallback heading: all caps and short
                letters = ''.join(ch for ch in pt if ch.isalpha())
                if (not is_heading) and letters and letters.isupper() and len(pt) <= 60:
                    is_heading = True
                if is_heading:
                    htxt = norm_heading(pt)
                    hints['headings'].append({'text': htxt})
                    current_section = htxt
                    if htxt not in sections_map:
                        sections_map[htxt] = {'name': htxt, 'subheadings': [], 'lines': [], 'bold_lines': []}
                    continue
                # Record paragraph under current section
                if current_section:
                    sections_map[current_section]['lines'].append(pt)
                # Detect bold subheadings
                has_bold = any((run.bold or False) for run in p.runs)
                if has_bold:
                    hints['bold_lines'].append(pt)
                    if current_section:
                        sections_map[current_section]['bold_lines'].append(pt)
                        # Use short bold lines as subheadings
                        if 2 <= len(pt) <= 100:
                            sections_map[current_section]['subheadings'].append(pt)
            # Education candidates (quick scan)
            for p in doc.paragraphs:
                t = p.text.strip()
                if t and ('college' in t.lower() or 'university' in t.lower()):
                    hints['education_candidates'].append(t)
            hints['sections'] = list(sections_map.values())
        elif file_type == 'pdf':
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            sections = []
            current = None
            sizes = []
            pages_data = []
            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                pd = page.get_text('dict')
                pages_data.append(pd)
                for b in pd.get('blocks', []):
                    for l in b.get('lines', []):
                        for s in l.get('spans', []):
                            sizes.append(s.get('size', 0))
            size_threshold = 0
            if sizes:
                # Consider headings as top 15% font sizes
                sizes_sorted = sorted(sizes)
                cutoff_index = int(0.85 * len(sizes_sorted))
                size_threshold = sizes_sorted[cutoff_index]
            def add_current():
                nonlocal current
                if current:
                    sections.append(current)
                    current = None
            for page_index, pd in enumerate(pages_data):
                for b in pd.get('blocks', []):
                    for l in b.get('lines', []):
                        line_text = ''.join(s.get('text', '') for s in l.get('spans', [])).strip()
                        if not line_text:
                            continue
                        # Heading if large font and mostly uppercase
                        max_size = max((s.get('size', 0) for s in l.get('spans', [])), default=0)
                        has_bold = any('Bold' in (s.get('font', '') or '') for s in l.get('spans', []))
                        letters = ''.join(ch for ch in line_text if ch.isalpha())
                        is_caps = letters.isupper() and len(letters) >= 3
                        if max_size >= size_threshold and is_caps:
                            add_current()
                            htxt = line_text.strip().replace(':', '').upper()
                            hints['headings'].append({'text': htxt})
                            current = {'name': htxt, 'subheadings': [], 'lines': [], 'bold_lines': []}
                        else:
                            # content line
                            if current:
                                current['lines'].append(line_text)
                                if has_bold or (is_caps and max_size >= (size_threshold * 0.9)):
                                    current['bold_lines'].append(line_text)
                                    if 2 <= len(line_text) <= 100:
                                        current['subheadings'].append(line_text)
            add_current()
            hints['sections'] = sections
            # Education candidates (simple)
            edu_cands = []
            for sec in sections:
                for ln in sec.get('lines', []):
                    if 'college' in ln.lower() or 'university' in ln.lower():
                        edu_cands.append(ln)
            hints['education_candidates'] = edu_cands
        else:
            # Images or other: provide just text and no hints
            pass
    except Exception as e:
        logger.warning(f"extract_text_with_hints_from_bytes: hints extraction fallback due to: {e}")
    return {'text': text, 'hints': hints}


if __name__ == "__main__":
    # Test the text extractor
    extractor = TextExtractor()
    
    # Example usage
    print("ResuMatch AI - Text Extractor")
    print("Supported formats: PDF, DOCX, JPG, PNG, BMP, TIFF, GIF")
    print("Usage: python extract_text.py <file_path>")
